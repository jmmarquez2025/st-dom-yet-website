#!/usr/bin/env python3
"""
Generate the St. Dominic Website Administrator Guide as a .docx file
from the Markdown source (docs/admin-guide.md).

Usage:
  python3 docs/generate-docx.py          # saves to ~/Desktop/
  npm run docs                            # same thing via npm

Requires: pip3 install python-docx
"""

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is required. Install it with:")
    print("  pip3 install python-docx")
    exit(1)

# ── Paths ──
SCRIPT_DIR = Path(__file__).parent
MD_PATH = SCRIPT_DIR / "admin-guide.md"
OUTPUT_PATH = Path.home() / "Desktop" / "St-Dominic-Website-Admin-Guide.docx"

BURGUNDY = RGBColor(0x6B, 0x1D, 0x2A)
GOLD = RGBColor(0xC5, 0xA5, 0x5A)
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x33, 0x33, 0x33)


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.color.rgb = BURGUNDY
        h.font.name = 'Calibri'
        if level == 1:
            h.font.size = Pt(22)
            h.paragraph_format.space_before = Pt(24)
        elif level == 2:
            h.font.size = Pt(16)
            h.paragraph_format.space_before = Pt(18)
        else:
            h.font.size = Pt(13)
            h.paragraph_format.space_before = Pt(12)


def add_title_page(doc):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ST. DOMINIC CATHOLIC CHURCH')
    run.font.size = Pt(28)
    run.font.color.rgb = BURGUNDY
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Website Administrator Guide')
    run.font.size = Pt(18)
    run.font.color.rgb = GRAY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Prepared for')
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Fr. Frassati Davis, O.P.  \u2022  '
        'Fr. Charles Marie Rooney, O.P.  \u2022  '
        'Parish Secretary'
    )
    run.font.size = Pt(12)
    run.bold = True

    doc.add_paragraph()

    for text in ['April 2026', 'Youngstown, Ohio']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY

    doc.add_page_break()


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx."""
    headers = [c.strip() for c in lines[start_idx].strip('|').split('|')]
    row_idx = start_idx + 2
    rows = []
    while row_idx < len(lines) and '|' in lines[row_idx] and lines[row_idx].strip():
        cells = [c.strip() for c in lines[row_idx].strip('|').split('|')]
        rows.append(cells)
        row_idx += 1
    return headers, rows, row_idx


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx < len(table.columns):
                table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()


def add_rich_paragraph(doc, text):
    """Add a paragraph with basic markdown bold/code formatting."""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif '`' in part:
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = p.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x88, 0x44, 0x44)
                elif cp:
                    p.add_run(cp)
        elif part:
            p.add_run(part)
    return p


def convert_md_to_docx(md_text, doc):
    lines = md_text.split('\n')
    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rules / page breaks
        if stripped == '---':
            doc.add_page_break()
            i += 1
            continue

        # Headings
        if stripped.startswith('####'):
            doc.add_heading(stripped.lstrip('#').strip(), level=3)
            i += 1
            continue
        if stripped.startswith('###'):
            doc.add_heading(stripped.lstrip('#').strip(), level=3)
            i += 1
            continue
        if stripped.startswith('##'):
            doc.add_heading(stripped.lstrip('#').strip(), level=2)
            i += 1
            continue
        if stripped.startswith('#'):
            doc.add_heading(stripped.lstrip('#').strip(), level=1)
            i += 1
            continue

        # Tables
        if '|' in stripped and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers, rows, end_idx = parse_table(lines, i)
            add_table(doc, headers, rows)
            i = end_idx
            continue

        # Blockquotes
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            if '**Tip:**' in text or '**TIP:**' in text:
                run = p.add_run('\u2728 TIP: ')
                run.bold = True
                run.font.color.rgb = GOLD
                clean = text.replace('**Tip:**', '').replace('**TIP:**', '').strip()
                run = p.add_run(clean)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            elif '**Important:**' in text or '**IMPORTANT:**' in text:
                run = p.add_run('\u26A0 IMPORTANT: ')
                run.bold = True
                run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                clean = text.replace('**Important:**', '').replace('**IMPORTANT:**', '').strip()
                run = p.add_run(clean)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                run = p.add_run(clean)
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+\.', stripped):
            text = re.sub(r'^\d+\.\s*', '', stripped)
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            doc.add_paragraph(clean, style='List Number')
            i += 1
            continue

        # Bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            doc.add_paragraph(clean, style='List Bullet')
            i += 1
            continue

        # Italic closing line
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.strip('*'))
            run.italic = True
            run.font.color.rgb = GRAY
            run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph
        add_rich_paragraph(doc, stripped)
        i += 1


def add_closing_page(doc):
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\u2720')
    run.font.size = Pt(36)
    run.font.color.rgb = BURGUNDY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Veritas')
    run.font.size = Pt(24)
    run.font.color.rgb = BURGUNDY
    run.italic = True

    doc.add_paragraph()
    for text in [
        'St. Dominic Catholic Church',
        '77 East Lucius Avenue, Youngstown, OH 44507',
        '(330) 783-1900  |  office@saintdominic.org',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY


def main():
    print(f"Reading {MD_PATH}...")
    md_text = MD_PATH.read_text(encoding='utf-8')

    doc = Document()
    setup_styles(doc)
    add_title_page(doc)
    convert_md_to_docx(md_text, doc)
    add_closing_page(doc)

    doc.save(str(OUTPUT_PATH))
    print(f"Saved to {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
